#the first step to read the PDF and extract the text into the document
import PyPDF2 
from docx import Document
from docx.shared import Inches
from PyPDF2 import PdfReader

def read_From_Pdf2(input_Pdf_File,output_docx_file):
# Open the PDF file and read its content
    print('in read_From_Pdf2 function')
    reader = open(input_Pdf_File, 'rb')
    pdf_reader = PyPDF2.PdfReader(reader)
    num_pages = len(pdf_reader.pages)
    content = ""
    
# Iterate through each page and extract the text
    for i in range(num_pages):
        page = pdf_reader.pages[i]
        content += page.extract_text()
        #print(page.extract_text())
    sentences = content.split('.')
# Create a new Word document and add the extracted text to it
    doc = Document()
    for sentence in sentences:
        if sentence.strip() != "":
            doc.add_paragraph(sentence.strip())
    doc.save(output_docx_file)
    # Close the PDF file
    reader.close()
    print('fin read_From_Pdf2 function')
    
if __name__ == "__main__":
    input_Pdf_File = r"D:\Doctoral thesis\Doctoral thesis database\DATA for THESIS\G7 database\countries\Germany\SAP\SAP CSR 2016.pdf"
    output_docx_file = r"D:\Doctoral thesis\Doctoral thesis database\DATA for THESIS\G7 database\countries\Germany\SAP\SAP 2016.docx"
    read_From_Pdf2(input_Pdf_File, output_docx_file)
    print('Pdf read successfully and extracted into docx')
	###############################################################################
#the second step for the assessment of the fog_index and pasting the value into excel sheet
import docx
from docx import Document
import openpyxl
from textblob import TextBlob
from nltk.corpus import stopwords
import nltk
import pyphen
nltk.download('punkt')

def calculate_fog_index(output_docx_file):
    # Open the Word document
    #doc = open(input_docx_file)
    doc = docx.Document(output_docx_file)
    text = []  # Initialize an empty list to store paragraphs
    
    # Extract the text from the document
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    # Join the list of text paragraphs into a single string
    text = '\n'.join(text)
    words = TextBlob(text).words

    # Segment the text into sentences
    sentences = TextBlob(text).sentences

    # Calculate the average words per sentence
    avg_words_per_sentence = len(words) / len(sentences)

    # Remove stopwords
    stop_words = set(stopwords.words('english'))
    filtered_words = [word for word in words if word.lower() not in stop_words]

    # Initialize Pyphen for syllable counting
    dic = pyphen.Pyphen(lang='en')

    # Count syllables for each word
    syllable_count = [len(dic.inserted(word).split('-')) for word in filtered_words]

    # Define criteria for complex words (e.g., words with more than two syllables)
    complex_word_count = sum(1 for count in syllable_count if count > 3)

    # Calculate the percentage of complex words
    percent_complex = (complex_word_count / len(filtered_words)) * 100
    return percent_complex

    # Calculate the Fog Index
    fog_index = 0.4 * (avg_words_per_sentence + percent_complex)

    print('fin calculate_fog_index function', fog_index)
    return fog_index

##############################################################


def insert_value_into_excel(file_path, sheet_name, company_name, year, value):
    # Load the workbook
    print('in insert_value_into_excel function')
    workbook = openpyxl.load_workbook(file_path)
    
    # Select the desired sheet
    sheet = workbook[sheet_name]
    # Find the row number based on the company name and year
    i=0
    for row in sheet.iter_rows(min_row=1, max_col=40, max_row=sheet.max_row, values_only=True):
        if ((row[0] == company_name) and (str(row[1]) == str(year))):
            # Update the value in the third column
            sheet.cell(row=i+1, column=36, value=value)
            break
        i=i+1
    else:
        # If the combination is not found, you can handle it accordingly
        print(f"No entry found for Company '{company_name}' in year '{year}'.")
    print('write in excel done')
    # Save the changes
    workbook.save(file_path)
    
if __name__ == "__main__":
    file_path = r"D:\Doctoral thesis\Doctoral thesis database\DATA for THESIS\G7 database\countries\TestT.xlsx"
    sheet_name = 'regression'
    company_name_to_insert = 'SAP'
    year_to_insert = 2016
    output_docx_file = r"D:\Doctoral thesis\Doctoral thesis database\DATA for THESIS\G7 database\countries\Germany\SAP\SAP 2016.docx"
    fog = calculate_fog_index(output_docx_file)
    print('fog index ', fog)
    insert_value_into_excel(file_path, sheet_name, company_name_to_insert, year_to_insert, fog)
    print('fin')